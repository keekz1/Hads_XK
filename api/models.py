from django.db import models
from django.contrib.auth.models import User
from django.utils import timezone
from django.db.models.signals import post_save
from django.dispatch import receiver
from decimal import Decimal
import os
# === USER PROFILE ===
class UserProfile(models.Model):
    TIER_CHOICES = [
        ('free', 'Free Tier (50 requests/day)'),
        ('premium', 'Premium ($19.99/month)'),
        ('unlimited', 'Unlimited ($49.99/month)'),
    ]
    
    PROVIDER_CHOICES = [
        ('groq', '🔥 Groq (FREE - 5M tokens/month)'),
        ('openai', 'OpenAI (GPT-3.5 Turbo)'),
    ]
    
    MODEL_CHOICES = [
        # Groq models
        ('llama-3.1-8b-instant', 'Llama 3.1 8B (Instant) - FREE'),
        ('mixtral-8x7b-32768', 'Mixtral 8x7B - FREE'),
        # OpenAI models
        ('gpt-3.5-turbo', 'GPT-3.5 Turbo'),
        ('gpt-3.5-turbo-instruct', 'GPT-3.5 Turbo Instruct'),
    ]
    
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    bio = models.TextField(blank=True)
    level = models.CharField(max_length=100, default="beginner")
    created_at = models.DateTimeField(auto_now_add=True)
    
    # Subscription info (users pay YOU)
    subscription_tier = models.CharField(
        max_length=20, 
        choices=TIER_CHOICES, 
        default='free'
    )
    
    # USER'S API KEYS (THEY provide these)
    groq_api_key = models.CharField(max_length=255, blank=True, null=True, help_text="FREE - 5M tokens/month at Groq.com")
    openai_api_key = models.CharField(max_length=255, blank=True, null=True, help_text="For GPT-3.5 Turbo")
    
    # Preferred AI provider (DEFAULT TO FREE GROQ!)
    preferred_provider = models.CharField(
        max_length=50,
        choices=PROVIDER_CHOICES,
        default='groq'  # Default to FREE GROQ!
    )
    
    # Preferred model
    preferred_model = models.CharField(
        max_length=100,
        choices=MODEL_CHOICES,
        default="llama-3.1-8b-instant",  # FREE Groq model
        help_text="Select model based on your provider"
    )
    
    # OpenAI specific info
    openai_key_type = models.CharField(
        max_length=20,
        choices=[
            ('unknown', 'Unknown'),
            ('free_tier', 'Free Tier'),
            ('pay_as_you_go', 'Pay-as-you-go'),
        ],
        default='unknown'
    )
    openai_key_credit_balance = models.DecimalField(max_digits=10, decimal_places=2, default=0.00, help_text="OpenAI credit balance if available") 
    openai_account_status = models.CharField(
        max_length=20,
        choices=[
            ('active', 'Active'),
            ('free_credits_expired', 'Free Credits Expired'),
            ('low_balance', 'Low Balance'),
            ('inactive', 'Inactive'),
        ],
        default='active',
        help_text="Status of the OpenAI account"
    )
    
    # Usage tracking (YOUR service limits)
    requests_today = models.IntegerField(default=0)
    tokens_this_month = models.IntegerField(default=0)
    pdf_analyses_this_month = models.IntegerField(default=0)
    data_analyses_this_month = models.IntegerField(default=0)
    
    # Monthly reset
    monthly_reset_date = models.DateField(default=timezone.now)
    
    # Rate limiting (YOUR service limits)
    requests_this_minute = models.IntegerField(default=0)
    last_request_time = models.DateTimeField(null=True, blank=True)
    last_reset_date = models.DateField(default=timezone.now)
    
    # Analytics
    last_active = models.DateTimeField(auto_now=True)
    
    # Stripe billing info (they pay YOU)
    stripe_customer_id = models.CharField(max_length=100, blank=True, null=True)
    stripe_subscription_id = models.CharField(max_length=100, blank=True, null=True)
    subscription_status = models.CharField(
        max_length=20,
        choices=[
            ('active', 'Active'),
            ('canceled', 'Canceled'),
            ('past_due', 'Past Due'),
            ('trialing', 'Trialing'),
        ],
        default='trialing'
    )
    
    # Current billing period
    current_period_start = models.DateField(null=True, blank=True)
    current_period_end = models.DateField(null=True, blank=True)
    
    # Reseller information
    reseller_code_used = models.CharField(max_length=50, blank=True, null=True, help_text="Reseller code used at signup")
    referred_by_code = models.CharField(max_length=50, blank=True, null=True, help_text="Referral code used at signup")
    
    def __str__(self):
        return f"{self.user.username}'s Profile"
    
    def has_api_key(self):
        """Check if user has API key for their preferred provider"""
        if self.preferred_provider == 'groq':
            return bool(self.groq_api_key)
        elif self.preferred_provider == 'openai':
            return bool(self.openai_api_key)
        return False
    
    def get_api_key(self):
        """Get the API key for the current preferred provider"""
        if self.preferred_provider == 'groq':
            return self.groq_api_key
        elif self.preferred_provider == 'openai':
            return self.openai_api_key
        return None
    
    def get_available_models(self):
        """Get available models for the current provider"""
        if self.preferred_provider == 'groq':
            return [
                ('llama-3.1-8b-instant', 'Llama 3.1 8B (Instant) - FREE'),
                ('mixtral-8x7b-32768', 'Mixtral 8x7B - FREE'),
            ]
        elif self.preferred_provider == 'openai':
            return [
                ('gpt-3.5-turbo', 'GPT-3.5 Turbo'),
                ('gpt-3.5-turbo-instruct', 'GPT-3.5 Turbo Instruct'),
            ]
        return []
    
    def validate_model_selection(self, model):
        """Validate that the selected model is compatible with the provider"""
        if self.preferred_provider == 'groq':
            return model in ['llama-3.1-8b-instant', 'mixtral-8x7b-32768']
        elif self.preferred_provider == 'openai':
            return model in ['gpt-3.5-turbo', 'gpt-3.5-turbo-instruct']
        return False
    
     
    
    def can_send_message(self):
        """Check if user can send a message (YOUR service limits)"""
        limits = self.get_tier_limits()
        
        # Check daily limit
        if self.requests_today >= limits['daily_requests']:
            return False, f"Daily limit reached ({self.requests_today}/{limits['daily_requests']})"
        
        # Check minute rate limit
        current_time = timezone.now()
        if self.last_request_time:
            time_diff = (current_time - self.last_request_time).total_seconds()
            if time_diff < 60 and self.requests_this_minute >= 5:
                return False, "Rate limit: 5 requests per minute. Please wait."
        
        return True, "OK"
    
    def record_request(self, tokens=0, is_pdf=False, is_data_analysis=False):
        """Record a user request (YOUR service tracking)"""
        current_time = timezone.now()
        
        # Reset daily counter if new day
        if current_time.date() != self.last_reset_date:
            self.requests_today = 0
            self.requests_this_minute = 0
            self.last_reset_date = current_time.date()
        
        # Reset minute counter if new minute
        if (not self.last_request_time or 
            (current_time - self.last_request_time).total_seconds() >= 60):
            self.requests_this_minute = 0
        
        # Update counters
        self.requests_today += 1
        self.requests_this_minute += 1
        self.tokens_this_month += tokens
        self.last_request_time = current_time
        
        if is_pdf:
            self.pdf_analyses_this_month += 1
        if is_data_analysis:
            self.data_analyses_this_month += 1
        
        self.save()
    
    def get_provider_info(self):
        """Get information about the current provider"""
        if self.preferred_provider == 'groq':
            return {
                'name': 'Groq',
                'is_free': True,
                'free_tokens': '5,000,000 tokens/month (FREE!)',
                'link': 'https://console.groq.com/keys',
                'models': self.get_available_models(),
            }
        elif self.preferred_provider == 'openai':
            return {
                'name': 'OpenAI',
                'is_free': self.openai_key_type == 'free_tier',
                'free_tokens': 'Free tier: Limited GPT-3.5 access' if self.openai_key_type == 'free_tier' else 'Pay-as-you-go',
                'link': 'https://platform.openai.com/api-keys',
                'models': self.get_available_models(),
                'credit_balance': float(self.openai_key_credit_balance),
                'account_status': self.openai_account_status,
            }
        return {}
    
    def update_openai_info(self, test_result):
        """Update OpenAI information from test results"""
        if 'key_type' in test_result:
            self.openai_key_type = test_result.get('key_type', 'unknown')
            self.openai_key_credit_balance = Decimal(str(test_result.get('credit_balance', 0)))
            self.openai_account_status = test_result.get('account_status', 'active')
            self.save()
    
    def is_subscription_active(self):
        """Check if user has an active subscription with YOU"""
        return self.subscription_status == 'active'
    
    def get_plan_price(self):
        """Get the price of user's current plan with YOU"""
        if self.subscription_tier == 'premium':
            return Decimal('19.99')
        elif self.subscription_tier == 'unlimited':
            return Decimal('49.99')
        return Decimal('0.00')
    
# In models.py - make sure this method exists exactly like this:
    def get_your_profit(self):
     """Calculate YOUR profit from this user"""
     plan_price = self.get_plan_price()
     if plan_price == Decimal('0.00'):
         return {'profit': Decimal('0.00'), 'margin_percent': 0}
    
    # Calculate YOUR costs
     stripe_fee = plan_price * Decimal('0.029') + Decimal('0.30')
     hosting_cost = Decimal('0.10')
    
    # YOUR profit
     your_profit = plan_price - stripe_fee - hosting_cost
     margin_percent = (your_profit / plan_price) * 100 if plan_price > 0 else 0
    
     return {
        'profit': your_profit,
        'margin_percent': round(float(margin_percent), 1)
    }
    
    def update_model_access_info(self, test_result):
        """Update model access information from test results"""
        if self.preferred_provider == 'openai' and 'key_type' in test_result:
            self.openai_key_type = test_result.get('key_type', 'unknown')
            self.openai_key_credit_balance = Decimal(str(test_result.get('credit_balance', 0)))
            self.openai_has_gpt4_access = test_result.get('has_gpt4_access', False)
            self.openai_has_gpt4_vision = test_result.get('has_gpt4_vision', False)
            self.openai_has_finetuning_access = test_result.get('has_finetuning_access', False)
            self.openai_total_models_count = test_result.get('total_models_count', 0)
            self.openai_account_status = test_result.get('account_status', 'active')
            self.openai_has_payment_method = test_result.get('has_payment_method', False)
            self.openai_can_use_paid_models = test_result.get('can_use_paid_models', False)
            self.openai_recommended_model = test_result.get('recommended_model', 'gpt-3.5-turbo')
            
            # Store available models
            gpt4_models = test_result.get('available_gpt4_models', [])
            gpt35_models = test_result.get('available_gpt35_models', [])
            
            self.openai_available_gpt4_models = ','.join(gpt4_models[:20])  # Limit to 20 models
            self.openai_available_gpt35_models = ','.join(gpt35_models[:20])
            
            self.last_model_detection = timezone.now()
            self.save()
            
        elif self.preferred_provider == 'groq' and 'models_count' in test_result:
            self.groq_total_models_count = test_result.get('models_count', 0)
            # You might want to store the actual model names if available
            if 'available_models' in test_result:
                models = test_result.get('available_models', [])
                self.groq_available_models = ','.join(models[:20])
            self.last_model_detection = timezone.now()
            self.save()
    
    def get_model_access_summary(self):
        """Get a summary of model access for the current provider"""
        if self.preferred_provider == 'openai':
            return {
                'provider': 'OpenAI',
                'account_type': self.openai_key_type,
                'credit_balance': float(self.openai_key_credit_balance),
                'has_gpt4_access': self.openai_has_gpt4_access,
                'has_gpt4_vision': self.openai_has_gpt4_vision,
                'has_finetuning_access': self.openai_has_finetuning_access,
                'total_models': self.openai_total_models_count,
                'account_status': self.openai_account_status,
                'has_payment_method': self.openai_has_payment_method,
                'can_use_paid_models': self.openai_can_use_paid_models,
                'recommended_model': self.openai_recommended_model,
                'available_gpt4_models': self.get_gpt4_models(),
                'available_gpt35_models': self.get_gpt35_models(),
                'last_detected': self.last_model_detection.isoformat() if self.last_model_detection else None,
            }
        elif self.preferred_provider == 'groq':
            return {
                'provider': 'Groq',
                'is_free': True,
                'total_models': self.groq_total_models_count,
                'available_models': self.get_available_models(),
                'free_tokens': '5,000,000 tokens/month',
                'last_detected': self.last_model_detection.isoformat() if self.last_model_detection else None,
            }
        else:
            return {
                'provider': self.preferred_provider.capitalize(),
                'available_models': self.get_available_models(),
            }
    
    def get_suggested_models(self):
        """Get suggested models based on account type"""
        if self.preferred_provider == 'openai':
            if self.openai_key_type == 'free_tier':
                return self.get_gpt35_models()
            elif self.openai_has_gpt4_access:
                # Return all available models
                return self.get_available_models()
            else:
                return self.get_gpt35_models()
        elif self.preferred_provider == 'groq':
            return self.get_available_models()
        else:
            return self.get_available_models()



    def get_tier_limits(self):
        """Get limits for user's tier (YOUR service limits)"""
        limits = {
            'free': {
                'daily_requests': 50,
                'can_analyze_pdf': False,
                'pdf_limit': 0,
                'data_analysis_limit': 0,
                'max_tokens_per_request': 2000,
                'support_level': 'basic',
            },
            'premium': {
                'daily_requests': 1000,
                'can_analyze_pdf': True,
                'pdf_limit': 5,
                'data_analysis_limit': 2,
                'max_tokens_per_request': 4000,
                'support_level': 'priority',
            },
            'unlimited': {
                'daily_requests': 5000,
                'can_analyze_pdf': True,
                'pdf_limit': 20,
                'data_analysis_limit': 10,
                'max_tokens_per_request': 8000,
                'support_level': 'unlimited',
            }
        }
        return limits.get(self.subscription_tier, limits['free'])
# Auto-create profile
@receiver(post_save, sender=User)
def create_user_profile(sender, instance, created, **kwargs):
    if created:
        profile = UserProfile.objects.create(user=instance)
        profile.save()

@receiver(post_save, sender=User)
def save_user_profile(sender, instance, **kwargs):
    if hasattr(instance, 'userprofile'):
        instance.userprofile.save()


class Subject(models.Model):
    name = models.CharField(max_length=100)
    description = models.TextField()
    
    # Analytics
    average_tokens_per_query = models.IntegerField(default=500)
    total_queries = models.IntegerField(default=0)

    def __str__(self):
        return self.name
   

class StudySession(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    subject = models.ForeignKey(Subject, on_delete=models.CASCADE)
    duration_minutes = models.IntegerField()
    notes = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    
    # AI usage
    ai_assistance_used = models.BooleanField(default=False)
    estimated_user_token_cost = models.DecimalField(max_digits=8, decimal_places=4, default=0.00)

    def __str__(self):
        return f"{self.user.username} - {self.subject.name}"


class StudyContent(models.Model):
    DIFFICULTY_CHOICES = [
        ('easy', 'Easy'),
        ('medium', 'Medium'),
        ('hard', 'Hard'),
    ]

    user = models.ForeignKey(User, on_delete=models.CASCADE)
    title = models.CharField(max_length=200)
    content = models.TextField()
    subject = models.CharField(max_length=100)
    difficulty = models.CharField(max_length=20, choices=DIFFICULTY_CHOICES)
    created_at = models.DateTimeField(auto_now_add=True)
    
    # AI generation tracking
    ai_generated = models.BooleanField(default=False)
    ai_model_used = models.CharField(max_length=100, blank=True, null=True)
    estimated_user_token_cost = models.DecimalField(max_digits=8, decimal_places=4, default=0.00)

    def __str__(self):
        return self.title


class AIConversation(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    prompt = models.TextField()
    response = models.TextField()
    subject = models.CharField(max_length=100)
    difficulty = models.CharField(max_length=50)
    created_at = models.DateTimeField(auto_now_add=True)
    
    # Add this new field:
    key_source = models.CharField(max_length=50, default='user', choices=[
        ('user', 'User API Key'),
        ('env', 'Environment API Key (Fallback)'),
        ('system', 'System Fallback Key')
    ], help_text="Source of the API key used for this request")
    
    # User's cost tracking
    model_used = models.CharField(max_length=100, default="llama-3.1-8b-instant")
    input_tokens = models.IntegerField(default=0)
    output_tokens = models.IntegerField(default=0)
    total_tokens = models.IntegerField(default=0)
    is_document_analysis = models.BooleanField(default=False)
    document_analyzed = models.ForeignKey(
        'UploadedDocument', 
        on_delete=models.SET_NULL, 
        null=True, 
        blank=True,
        related_name='analyses'
    )
    # Estimated cost to USER (0 for FREE Groq!)
    estimated_user_cost = models.DecimalField(max_digits=8, decimal_places=6, default=0.0000)
    
    # YOUR service fee (tiny profit per request)
    your_service_fee = models.DecimalField(max_digits=8, decimal_places=6, default=0.0001)
    
    response_time_ms = models.IntegerField(default=0)
    
    # Which API provider was used
    api_provider = models.CharField(max_length=50, default="groq")
    
    # Add this database field
    is_free = models.BooleanField(default=False, help_text="Whether this request was FREE (Groq or free tier OpenAI)")
    
    # Caching info
    was_cached = models.BooleanField(default=False)
    cache_hit = models.BooleanField(default=False)
    
    # User's tier at time of request
    user_tier_at_time = models.CharField(max_length=20, default='free')
 
    def __str__(self):
        return f"{self.user.username} - {self.subject}"
    
    def save(self, *args, **kwargs):
        # Auto-calculate total tokens
        if self.input_tokens and self.output_tokens:
            self.total_tokens = self.input_tokens + self.output_tokens  # FIXED: removed ?//
        
        # Auto-calculate is_free based on provider
        if self.api_provider == 'groq':
            self.is_free = True
        elif self.api_provider == 'openai':
            # Check if it's a free tier request (only GPT-3.5 models are free)
            if 'gpt-3.5' in self.model_used.lower():
                self.is_free = True
            else:
                self.is_free = False
        else:
            self.is_free = False
        
        # Update user's profile usage (only if not cached)
        if self.user and not self.was_cached:
            try:
                profile = self.user.userprofile
                profile.record_request(tokens=self.total_tokens)
            except:
                pass
        
        super().save(*args, **kwargs)
    
    @property
    def your_profit(self):
        """YOUR profit from this single request"""
        return self.your_service_fee
 

# === BILLING & ANALYTICS MODELS ===

class SubscriptionPlan(models.Model):
    """Plans users pay YOU for"""
    TIER_CHOICES = [
        ('free', 'Free'),
        ('premium', 'Premium'),
        ('unlimited', 'Unlimited'),
    ]
    
    name = models.CharField(max_length=100)
    tier = models.CharField(max_length=20, choices=TIER_CHOICES, unique=True)
    monthly_price = models.DecimalField(max_digits=6, decimal_places=2)
    
    # Description for UI
    description = models.TextField(default="Get started with FREE Groq AI!")
    features = models.JSONField(default=list)
    
    # Stripe IDs
    stripe_price_id = models.CharField(max_length=100)
    
    # Suggested limits (for UI)
    suggested_daily_requests = models.IntegerField(default=50)
    suggested_monthly_tokens = models.IntegerField(default=250000)
    includes_pdf_analysis = models.BooleanField(default=False)
    includes_data_analysis = models.BooleanField(default=False)
    
    # Your profit metrics
    your_cost_per_user = models.DecimalField(max_digits=6, decimal_places=2, default=Decimal('0.10'))
    
    def __str__(self):
        return f"{self.name} - ${self.monthly_price}/month"
    
    @property
    def your_profit_per_user(self):
        """YOUR profit per user per month"""
        return self.monthly_price - self.your_cost_per_user
    
    @property
    def your_profit_margin(self):
        """YOUR profit margin percentage"""
        if self.monthly_price == 0:
            return 0
        return round((self.your_profit_per_user / self.monthly_price) * 100, 1)


class UserSubscription(models.Model):
    """Track subscriptions users have with YOU"""
    STATUS_CHOICES = [
        ('active', 'Active'),
        ('canceled', 'Canceled'),
        ('past_due', 'Past Due'),
        ('trialing', 'Trialing'),
    ]
    
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    plan = models.ForeignKey(SubscriptionPlan, on_delete=models.CASCADE)
    status = models.CharField(max_length=20, choices=STATUS_CHOICES)
    stripe_subscription_id = models.CharField(max_length=100)
    stripe_customer_id = models.CharField(max_length=100)
    current_period_start = models.DateTimeField()
    current_period_end = models.DateTimeField()
    canceled_at = models.DateTimeField(null=True, blank=True)
    
    # Trial info
    trial_start = models.DateTimeField(null=True, blank=True)
    trial_end = models.DateTimeField(null=True, blank=True)
    
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    def __str__(self):
        return f"{self.user.username} - {self.plan.name} - {self.status}"
    
    @property
    def is_active(self):
        return self.status == 'active' or (self.status == 'trialing' and self.trial_end > timezone.now())
    
    @property
    def your_profit_this_period(self):
        """Calculate YOUR profit from this subscription"""
        if self.status != 'active':
            return Decimal('0.00')
        
        # Your revenue
        revenue = self.plan.monthly_price
        
        # Your costs
        stripe_fee = revenue * Decimal('0.029') + Decimal('0.30')  # Stripe fees
        hosting_cost = Decimal('0.10')  # Per user hosting
        
        total_cost = stripe_fee + hosting_cost
        profit = revenue - total_cost
        
        return {
            'revenue': revenue,
            'costs': total_cost,
            'profit': profit,
            'margin': round((profit / revenue) * 100, 1) if revenue > 0 else 0
        }


class BillingTransaction(models.Model):
    """Track payments from users to YOU"""
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    timestamp = models.DateTimeField(auto_now_add=True)
    
    # Payment details
    amount = models.DecimalField(max_digits=6, decimal_places=2)
    currency = models.CharField(max_length=3, default="USD")
    
    # What they paid for
    plan = models.ForeignKey(SubscriptionPlan, on_delete=models.SET_NULL, null=True)
    period_start = models.DateField()
    period_end = models.DateField()
    
    # YOUR costs (minimal)
    stripe_fee = models.DecimalField(max_digits=6, decimal_places=2)
    hosting_fee = models.DecimalField(max_digits=6, decimal_places=2, default=Decimal('0.10'))
    
    # Stripe info
    stripe_payment_intent_id = models.CharField(max_length=100)
    stripe_invoice_id = models.CharField(max_length=100, blank=True, null=True)
    stripe_invoice_url = models.URLField(blank=True, null=True)
    
    # Status
    status = models.CharField(
        max_length=20,
        choices=[
            ('succeeded', 'Succeeded'),
            ('failed', 'Failed'),
            ('refunded', 'Refunded'),
        ],
        default='succeeded'
    )
    
    class Meta:
        indexes = [
            models.Index(fields=['user', 'timestamp']),
            models.Index(fields=['status', 'timestamp']),
        ]
    
    def __str__(self):
        return f"{self.user.username} - ${self.amount} - {self.timestamp.date()}"
    
    @property
    def your_profit(self):
        """Calculate YOUR pure profit from this transaction"""
        if self.status != 'succeeded':
            return Decimal('0.00')
        
        return self.amount - self.stripe_fee - self.hosting_fee
    
    @property
    def profit_margin(self):
        """Calculate YOUR profit margin percentage"""
        if self.amount == 0 or self.status != 'succeeded':
            return 0
        return round((self.your_profit / self.amount) * 100, 1)


class APIProxyLog(models.Model):
    """Log all API calls through YOUR proxy"""
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    timestamp = models.DateTimeField(auto_now_add=True)
    
    # Request details
    endpoint = models.CharField(max_length=100)
    model = models.CharField(max_length=100)
    input_tokens = models.IntegerField(default=0)
    output_tokens = models.IntegerField(default=0)
    total_tokens = models.IntegerField(default=0)
    document = models.ForeignKey(
        'UploadedDocument',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='api_logs'
    )
    # Cost to USER (0 for FREE Groq!)
    estimated_user_cost = models.DecimalField(max_digits=8, decimal_places=6, default=0.0000)
    
    # YOUR service fee (profit)
    your_service_fee = models.DecimalField(max_digits=8, decimal_places=6, default=0.0001)
    
    # Response details
    response_time_ms = models.IntegerField(default=0)
    success = models.BooleanField(default=True)
    error_message = models.TextField(blank=True, null=True)
    
    # Which provider was used
    provider = models.CharField(max_length=50, default="groq")
    
    # Request metadata
    request_type = models.CharField(
        max_length=50,
        choices=[
            ('chat', 'Chat'),
            ('pdf', 'PDF Analysis'),
            ('data', 'Data Analysis'),
            ('other', 'Other'),
        ],
        default='chat'
    )
    
    class Meta:
        indexes = [
            models.Index(fields=['user', 'timestamp']),
            models.Index(fields=['timestamp']),
            models.Index(fields=['provider', 'timestamp']),
        ]
    
    def __str__(self):
        return f"{self.user.username} - {self.provider} - ${self.your_service_fee}"
    
    @property
    def your_profit(self):
        """YOUR profit from this request"""
        return self.your_service_fee
    
    @property
    def is_free(self):
        """Check if this request was FREE (Groq)"""
        return self.provider == 'groq'


class APIAutoSetupRequest(models.Model):
    """Track automated API account creation for users"""
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    provider = models.CharField(
        max_length=50,
        choices=[
            ('groq', '🔥 Groq (FREE)'),
            ('openai', 'OpenAI'),  # Removed anthropic and gemini
        ],
        default='groq'
    )
    
    # Setup details
    requested_at = models.DateTimeField(auto_now_add=True)
    completed_at = models.DateTimeField(null=True, blank=True)
    
    # Status tracking
    STATUS_CHOICES = [
        ('pending', 'Pending'),
        ('processing', 'Processing'),
        ('completed', 'Completed'),
        ('failed', 'Failed'),
    ]
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='pending')
    
    # Result data
    api_key = models.CharField(max_length=255, blank=True, null=True)
    account_email = models.EmailField(blank=True, null=True)
    account_id = models.CharField(max_length=100, blank=True, null=True)
    
    # Error info
    error_message = models.TextField(blank=True, null=True)
    error_code = models.CharField(max_length=50, blank=True, null=True)
    
    # Instructions for user
    setup_instructions = models.TextField(blank=True, null=True)
    
    # FREE tokens info (for Groq)
    free_tokens_info = models.CharField(max_length=100, blank=True, null=True, default="5,000,000 tokens/month")
    
    def __str__(self):
        return f"{self.user.username} - {self.provider} - {self.status}"
    
    @property
    def is_completed(self):
        return self.status == 'completed'
    
    @property
    def took_seconds(self):
        if self.completed_at:
            return (self.completed_at - self.requested_at).total_seconds()
        return None
    
    @property
    def is_free(self):
        """Check if this provider is FREE"""
        return self.provider == 'groq'


class RateLimitLog(models.Model):
    """Track rate limit hits on YOUR service"""
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    timestamp = models.DateTimeField(auto_now_add=True)
    limit_type = models.CharField(max_length=50)
    limit_value = models.IntegerField()
    current_usage = models.IntegerField()
    request_path = models.CharField(max_length=200, blank=True, null=True)
    
    def __str__(self):
        return f"{self.user.username} - {self.limit_type} limit hit"


class ProfitAnalytics(models.Model):
    """Track YOUR overall profits"""
    date = models.DateField(unique=True)
    
    # Revenue from users
    total_revenue = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    
    # Your costs
    stripe_fees = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    hosting_costs = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    
    # User counts
    active_free_users = models.IntegerField(default=0)
    active_premium_users = models.IntegerField(default=0)
    active_unlimited_users = models.IntegerField(default=0)
    
    # Usage stats
    total_requests = models.IntegerField(default=0)
    total_tokens_proxied = models.BigIntegerField(default=0)
    total_your_service_fees = models.DecimalField(max_digits=10, decimal_places=4, default=0)
    
    # FREE Groq usage stats
    groq_requests = models.IntegerField(default=0)
    free_tokens_used = models.BigIntegerField(default=0)
    
    # Reseller stats
    total_reseller_commissions = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    total_reseller_payouts = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    
    # Calculated fields
    total_costs = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    net_profit = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    profit_margin = models.DecimalField(max_digits=5, decimal_places=2, default=0)
    
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name_plural = "Profit Analytics"
        ordering = ['-date']
    
    def __str__(self):
        return f"Profit Analytics - {self.date} - ${self.net_profit}"
    
    def save(self, *args, **kwargs):
        # Auto-calculate
        self.total_costs = self.stripe_fees + self.hosting_costs
        self.net_profit = self.total_revenue - self.total_costs - self.total_reseller_commissions
        
        if self.total_revenue > 0:
            self.profit_margin = (self.net_profit / self.total_revenue) * 100
        
        super().save(*args, **kwargs)
    
    @classmethod
    def update_daily_analytics(cls):
        """Update daily profit analytics"""
        today = timezone.now().date()
        
        # Get all billing transactions for today
        today_transactions = BillingTransaction.objects.filter(
            timestamp__date=today,
            status='succeeded'
        )
        
        # Calculate revenue
        total_revenue = sum(t.amount for t in today_transactions)
        stripe_fees = sum(t.stripe_fee for t in today_transactions)
        
        # Get user counts
        profiles = UserProfile.objects.all()
        active_free = profiles.filter(subscription_tier='free', subscription_status='active').count()
        active_premium = profiles.filter(subscription_tier='premium', subscription_status='active').count()
        active_unlimited = profiles.filter(subscription_tier='unlimited', subscription_status='active').count()
        
        # Get usage stats
        proxy_logs = APIProxyLog.objects.filter(timestamp__date=today)
        total_requests = proxy_logs.count()
        total_tokens = sum(log.total_tokens for log in proxy_logs)
        total_service_fees = sum(log.your_service_fee for log in proxy_logs)
        
        # Get FREE Groq usage
        groq_logs = proxy_logs.filter(provider='groq')
        groq_requests = groq_logs.count()
        free_tokens = sum(log.total_tokens for log in groq_logs)
        
        # Get reseller commissions
        today_commissions = ResellerCommission.objects.filter(created_at__date=today, status='paid')
        total_commissions = sum(c.commission_amount for c in today_commissions)
        
        # Get reseller payouts
        today_payouts = ResellerPayout.objects.filter(created_at__date=today, status='completed')
        total_payouts = sum(p.amount for p in today_payouts)
        
        # Estimate hosting costs ($0.10 per active user per day)
        hosting_costs = (active_free + active_premium + active_unlimited) * Decimal('0.10')
        
        # Create or update analytics
        analytics, created = cls.objects.get_or_create(date=today)
        analytics.total_revenue = total_revenue
        analytics.stripe_fees = stripe_fees
        analytics.hosting_costs = hosting_costs
        analytics.active_free_users = active_free
        analytics.active_premium_users = active_premium
        analytics.active_unlimited_users = active_unlimited
        analytics.total_requests = total_requests
        analytics.total_tokens_proxied = total_tokens
        analytics.total_your_service_fees = total_service_fees
        analytics.groq_requests = groq_requests
        analytics.free_tokens_used = free_tokens
        analytics.total_reseller_commissions = total_commissions
        analytics.total_reseller_payouts = total_payouts
        analytics.save()
        
        return analytics


# === RESELLER MODELS ===

class Reseller(models.Model):
    """Users who resell the StudyPilot platform"""
    STATUS_CHOICES = [
        ('pending', 'Pending Approval'),
        ('active', 'Active'),
        ('suspended', 'Suspended'),
        ('rejected', 'Rejected'),
    ]
    
    user = models.OneToOneField(User, on_delete=models.CASCADE, related_name='reseller')
    name = models.CharField(max_length=100, help_text="Display name for reseller")
    company = models.CharField(max_length=100, blank=True, null=True)
    website = models.URLField(blank=True, null=True)
    description = models.TextField(blank=True, null=True)
    
    # Reseller settings
    code = models.CharField(max_length=50, unique=True, help_text="Unique referral code")
    default_commission_rate = models.DecimalField(
        max_digits=5, 
        decimal_places=2, 
        default=Decimal('0.30'),  # 30% default commission
        help_text="Default commission rate (0.00-1.00)"
    )
    discount_percent = models.DecimalField(
        max_digits=5, 
        decimal_places=2, 
        default=Decimal('0.00'),
        help_text="Discount percentage to offer clients (0-100)"
    )
    signup_bonus = models.TextField(blank=True, null=True, help_text="Special offer for signups")
    
    # Status
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='pending')
    is_active = models.BooleanField(default=False, help_text="Active resellers can accept new clients")
    
    # Approval tracking
    approved_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, blank=True, related_name='approved_resellers')
    approved_at = models.DateTimeField(null=True, blank=True)
    
    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    def __str__(self):
        return f"{self.name} ({self.code})"
    
    def calculate_stats(self):
        """Calculate reseller statistics"""
        clients = self.resellerclient_set.all()
        active_clients = clients.filter(status='active')
        
        # Calculate total earnings
        total_earnings = ResellerCommission.objects.filter(
            reseller=self,
            status='paid'
        ).aggregate(total=models.Sum('commission_amount'))['total'] or Decimal('0.00')
        
        # Calculate pending commissions
        pending_commissions = ResellerCommission.objects.filter(
            reseller=self,
            status='pending'
        ).aggregate(total=models.Sum('commission_amount'))['total'] or Decimal('0.00')
        
        # Calculate available balance (paid commissions not yet withdrawn)
        paid_not_withdrawn = ResellerCommission.objects.filter(
            reseller=self,
            status='paid',
            payout__isnull=True
        ).aggregate(total=models.Sum('commission_amount'))['total'] or Decimal('0.00')
        
        return {
            'total_clients': clients.count(),
            'active_clients': active_clients.count(),
            'total_earnings': total_earnings,
            'pending_commissions': pending_commissions,
            'available_balance': paid_not_withdrawn,
            'average_commission_rate': float(self.default_commission_rate),
        }
    
    def save(self, *args, **kwargs):
        if not self.code:
            # Generate unique code
            import uuid
            self.code = f"RS{uuid.uuid4().hex[:8].upper()}"
        
        if self.status == 'active' and not self.is_active:
            self.is_active = True
        elif self.status != 'active' and self.is_active:
            self.is_active = False
            
        super().save(*args, **kwargs)


class ResellerClient(models.Model):
    """Users who signed up through a reseller"""
    STATUS_CHOICES = [
        ('active', 'Active'),
        ('inactive', 'Inactive'),
        ('suspended', 'Suspended'),
    ]
    
    reseller = models.ForeignKey(Reseller, on_delete=models.CASCADE)
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='reseller_clients')
    
    # Commission settings
    commission_rate = models.DecimalField(
        max_digits=5, 
        decimal_places=2, 
        default=Decimal('0.30'),
        help_text="Commission rate for this client (0.00-1.00)"
    )
    
    # Status
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='active')
    
    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        unique_together = ['reseller', 'user']
    
    def __str__(self):
        return f"{self.user.username} -> {self.reseller.name}"
    
    def calculate_total_commission(self):
        """Calculate total commission earned from this client"""
        return ResellerCommission.objects.filter(
            reseller=self.reseller,
            client=self,
            status='paid'
        ).aggregate(total=models.Sum('commission_amount'))['total'] or Decimal('0.00')


class ResellerCommission(models.Model):
    """Commissions earned by resellers"""
    STATUS_CHOICES = [
        ('pending', 'Pending'),
        ('paid', 'Paid'),
        ('cancelled', 'Cancelled'),
    ]
    
    TRANSACTION_TYPES = [
        ('subscription', 'Subscription Payment'),
        ('service_fee', 'Service Fee'),
        ('pdf_analysis', 'PDF Analysis'),
        ('data_analysis', 'Data Analysis'),
        ('other', 'Other'),
    ]
    
    reseller = models.ForeignKey(Reseller, on_delete=models.CASCADE)
    client = models.ForeignKey(ResellerClient, on_delete=models.CASCADE, null=True, blank=True)
    
    # Transaction details
    transaction_type = models.CharField(max_length=20, choices=TRANSACTION_TYPES, default='service_fee')
    
    # Link to original transaction
    transaction = models.ForeignKey(APIProxyLog, on_delete=models.SET_NULL, null=True, blank=True, related_name='reseller_commissions')
    subscription_transaction = models.ForeignKey(BillingTransaction, on_delete=models.SET_NULL, null=True, blank=True, related_name='reseller_commissions')
    
    # Commission details
    commission_rate = models.DecimalField(max_digits=5, decimal_places=2)
    commission_amount = models.DecimalField(max_digits=10, decimal_places=4)
    
    # Payout tracking
    payout = models.ForeignKey('ResellerPayout', on_delete=models.SET_NULL, null=True, blank=True, related_name='commissions')
    
    # Status
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='pending')
    
    # Notes
    notes = models.TextField(blank=True, null=True)
    
    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    def __str__(self):
        return f"{self.reseller.name}: ${self.commission_amount} ({self.status})"
    
    def save(self, *args, **kwargs):
        # If linked to a transaction, get details automatically
        if self.transaction and not self.commission_amount:
            self.commission_amount = self.transaction.your_service_fee * self.commission_rate
            self.transaction_type = 'service_fee'
        
        if self.subscription_transaction and not self.commission_amount:
            self.commission_amount = self.subscription_transaction.amount * self.commission_rate
            self.transaction_type = 'subscription'
            
        super().save(*args, **kwargs)


class ResellerPayout(models.Model):
    """Payouts to resellers"""
    STATUS_CHOICES = [
        ('pending', 'Pending'),
        ('processing', 'Processing'),
        ('completed', 'Completed'),
        ('failed', 'Failed'),
        ('cancelled', 'Cancelled'),
    ]
    
    PAYOUT_METHODS = [
        ('stripe', 'Stripe'),
        ('paypal', 'PayPal'),
        ('bank_transfer', 'Bank Transfer'),
        ('crypto', 'Cryptocurrency'),
    ]
    
    reseller = models.ForeignKey(Reseller, on_delete=models.CASCADE, related_name='payouts')
    amount = models.DecimalField(max_digits=10, decimal_places=2)
    payout_method = models.CharField(max_length=20, choices=PAYOUT_METHODS, default='stripe')
    
    # Payment details
    payment_reference = models.CharField(max_length=100, blank=True, null=True)
    transaction_id = models.CharField(max_length=100, blank=True, null=True)
    
    # Status
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='pending')
    
    # Notes
    notes = models.TextField(blank=True, null=True)
    
    # Processing info
    processed_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, blank=True)
    processed_at = models.DateTimeField(null=True, blank=True)
    
    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    def __str__(self):
        return f"Payout #{self.id}: ${self.amount} to {self.reseller.name}"
    
    def save(self, *args, **kwargs):
        if self.status == 'completed' and not self.processed_at:
            self.processed_at = timezone.now()
        super().save(*args, **kwargs)
        
        
 
# api/models.py
class UploadedDocument(models.Model):
    DOCUMENT_TYPES = [
        ('pdf', 'PDF Document'),
        ('docx', 'Word Document'),
        ('txt', 'Text File'),
        ('image', 'Image File'),
        ('excel', 'Excel File'),
        ('ppt', 'PowerPoint File'),
    ]
    
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    file = models.FileField(upload_to='documents/%Y/%m/%d/')
    file_name = models.CharField(max_length=255)
    file_size = models.IntegerField(default=0)
    file_type = models.CharField(max_length=10, choices=DOCUMENT_TYPES)
    uploaded_at = models.DateTimeField(auto_now_add=True)
    extracted_text = models.TextField(blank=True, null=True)
    page_count = models.IntegerField(default=0)
    is_processed = models.BooleanField(default=False)
    
 
    
    # AI analysis metadata
    analyzed_at = models.DateTimeField(null=True, blank=True)
    analysis_model = models.CharField(max_length=100, blank=True, null=True)
    analysis_tokens = models.IntegerField(default=0)
    
    def __str__(self):
        return f"{self.user.username} - {self.file_name}"
    
    def save(self, *args, **kwargs):
        # Set file name and size
        if self.file:
            self.file_name = os.path.basename(self.file.name)
            try:
                self.file_size = self.file.size
            except:
                pass
            
            # Determine file type from extension
            ext = os.path.splitext(self.file_name)[1].lower()
            if ext == '.pdf':
                self.file_type = 'pdf'
            elif ext in ['.docx', '.doc']:
                self.file_type = 'docx'
            elif ext == '.txt':
                self.file_type = 'txt'
            elif ext in ['.jpg', '.jpeg', '.png', '.gif']:
                self.file_type = 'image'
            elif ext in ['.xlsx', '.xls', '.csv']:
                self.file_type = 'excel'
            elif ext in ['.ppt', '.pptx']:
                self.file_type = 'ppt'
        
        super().save(*args, **kwargs)
    
    def extract_text(self):
        """Extract text from the document"""
        try:
            text = ""
            
            if self.file_type == 'pdf':
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(self.file)
                self.page_count = len(pdf_reader.pages)
                for page_num in range(self.page_count):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
                
            elif self.file_type == 'docx':
                import docx
                doc = docx.Document(self.file)
                self.page_count = 1  # Approximate
                for para in doc.paragraphs:
                    text += para.text + "\n"
                    
            elif self.file_type == 'txt':
                self.file.seek(0)
                text = self.file.read().decode('utf-8', errors='ignore')
                self.page_count = 1
                
            elif self.file_type == 'image':
                # For images, we'll use OCR (optional)
                text = f"[Image file: {self.file_name}]"
                self.page_count = 1
                
            else:
                text = f"[{self.file_type.upper()} file: {self.file_name}]"
                self.page_count = 1
            
            self.extracted_text = text[:100000]  # Limit to 100k chars
            self.is_processed = True
            self.save()
            return text
            
        except Exception as e:
            print(f"Error extracting text: {e}")
            return f"Error processing file: {str(e)}"
    
    def get_summary(self, max_chars=500):
        """Get a summary of the document"""
        if not self.extracted_text:
            return "No text extracted from document"
        
        # Simple summary: first and last parts
        text = self.extracted_text
        if len(text) > max_chars:
            first_part = text[:max_chars//2]
            last_part = text[-max_chars//2:] if len(text) > max_chars else ""
            return f"{first_part}...\n\n...{last_part}"
        return text
    
    @property
    def file_size_mb(self):
        """Get file size in MB"""
        return round(self.file_size / (1024 * 1024), 2) if self.file_size else 0        